# -*- coding: utf-8 -*-
from docx import Document
from pathlib import Path

p = Path(r"G:\Mi unidad\Trabajos\Empleo\CUN\Cursos\Especializacion\PROYECTO I\Especializacion_En_Inteligencia_Artificial_Proyecto_I_ESP329.docx")
d = Document(str(p))
out = Path(r"G:\Mi unidad\Trabajos\Empleo\CUN\Cursos\config\slides\_tmp_esp329_extract.txt")
lines = []
lines.append("=== PARAS ===")
for i, para in enumerate(d.paragraphs):
    t = para.text.strip()
    if t:
        style = para.style.name if para.style else ""
        lines.append(f"[{i}|{style}] {t}")
lines.append("=== TABLES ===")
for ti, table in enumerate(d.tables):
    lines.append(f"--- TABLE {ti} ---")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
        lines.append(f"R{ri}: " + " || ".join(cells))
out.write_text("\n".join(lines), encoding="utf-8")
print(f"Wrote {out} ({len(lines)} lines)")
