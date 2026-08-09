# -*- coding: utf-8 -*-
"""Identidad visual CUN para documentos .docx (estudiantes y builds).

Fuente de marca: ``config/universidades/cun.json`` (colores + logo).
Usar desde ``guion_md_a_docx.convert(..., brand=True)`` o llamar
``apply_cun_brand_docx(doc, ...)`` sobre un Document ya armado.

Paleta:
  verde #007433 · lima #91DC00 · navy #0C2340
Logo: ``config/slides/assets/logo_cun_solo.png``
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
ASSETS = Path(__file__).resolve().parent / "assets"
LOGO_SOLO = ASSETS / "logo_cun_solo.png"
LOGO_FULL = ASSETS / "logo_cun.png"

# Colores canónicos (cun.json → marca.colores)
VERDE = "007433"
LIMA = "91DC00"
NAVY = "0C2340"
GRIS = "2B2B2B"
VERDE_CLARO = "E8F5EC"
ALT_ROW = "F5F7F5"

RGB_VERDE = RGBColor(0x00, 0x74, 0x33)
RGB_LIMA = RGBColor(0x91, 0xDC, 0x00)
RGB_NAVY = RGBColor(0x0C, 0x23, 0x40)
RGB_GRIS = RGBColor(0x2B, 0x2B, 0x2B)
RGB_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT = "Calibri"


def logo_path() -> str | None:
    for p in (LOGO_SOLO, LOGO_FULL):
        if p.is_file():
            return str(p)
    return None


def _shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_run_font(run, *, size=11, bold=False, color=RGB_GRIS, name=FONT) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), name)


def _add_header_band(doc: Document, subtitle: str | None = None) -> None:
    """Tabla superior: logo | CUN + subtítulo, con franja verde/lima."""
    # Línea lima decorativa
    tip = doc.add_paragraph()
    tip.paragraph_format.space_before = Pt(0)
    tip.paragraph_format.space_after = Pt(0)
    r = tip.add_run("━" * 48)
    _set_run_font(r, size=6, color=RGB_LIMA, bold=True)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    c0, c1 = table.cell(0, 0), table.cell(0, 1)
    _shade_cell(c0, VERDE_CLARO)
    _shade_cell(c1, VERDE_CLARO)

    # Logo
    c0.paragraphs[0].text = ""
    p_logo = c0.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo = logo_path()
    if logo:
        try:
            run = p_logo.add_run()
            run.add_picture(logo, width=Inches(1.15))
        except Exception:
            r = p_logo.add_run("CUN")
            _set_run_font(r, size=16, bold=True, color=RGB_VERDE)
    else:
        r = p_logo.add_run("CUN")
        _set_run_font(r, size=16, bold=True, color=RGB_VERDE)

    # Texto institucional
    c1.paragraphs[0].text = ""
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run("Corporación Unificada Nacional\nde Educación Superior")
    _set_run_font(r1, size=9, bold=True, color=RGB_NAVY)
    if subtitle:
        p2 = c1.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(subtitle)
        _set_run_font(r2, size=8, color=RGB_VERDE)

    # Barra navy + acento lima
    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(2)
    bar.paragraph_format.space_after = Pt(8)
    r = bar.add_run("━" * 20)
    _set_run_font(r, size=8, color=RGB_NAVY, bold=True)
    r2 = bar.add_run(" ━ ")
    _set_run_font(r2, size=8, color=RGB_LIMA, bold=True)
    r3 = bar.add_run("━" * 20)
    _set_run_font(r3, size=8, color=RGB_VERDE, bold=True)


def _set_section_footer(doc: Document, footer_text: str) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Línea lima arriba del pie
    r0 = p.add_run("― ")
    _set_run_font(r0, size=7, color=RGB_LIMA)
    r = p.add_run(footer_text)
    _set_run_font(r, size=8, color=RGB_NAVY)
    r2 = p.add_run(" ―")
    _set_run_font(r2, size=7, color=RGB_LIMA)

    # Márgenes cómodos
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def style_headings_cun(doc: Document) -> None:
    """Recolorea títulos ya insertados (H1 navy, H2 verde, resto navy)."""
    for p in doc.paragraphs:
        if not p.runs:
            continue
        # Heurística: párrafos con un solo run bold grande ≈ heading del convertidor
        sizes = [r.font.size.pt for r in p.runs if r.font.size]
        if not sizes:
            continue
        mx = max(sizes)
        all_bold = all(bool(r.bold) for r in p.runs if r.text.strip())
        if not all_bold:
            continue
        if mx >= 18:
            color = RGB_NAVY
        elif mx >= 14:
            color = RGB_VERDE
        elif mx >= 12:
            color = RGB_NAVY
        else:
            continue
        for r in p.runs:
            r.font.color.rgb = color
            r.font.name = FONT


def style_tables_cun(doc: Document) -> None:
    """Encabezados de tabla navy + filas alternas verdes suaves."""
    for table in doc.tables:
        if not table.rows:
            continue
        # Skip brand header table (2 cols, 1 row, verde claro)
        if len(table.rows) == 1 and len(table.columns) == 2:
            continue
        for j, cell in enumerate(table.rows[0].cells):
            _shade_cell(cell, NAVY)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGB_WHITE
                    r.bold = True
                    r.font.name = FONT
        for i, row in enumerate(table.rows[1:], start=1):
            if i % 2 == 0:
                for cell in row.cells:
                    _shade_cell(cell, ALT_ROW)


def apply_cun_brand_docx(
    doc: Document,
    *,
    subtitle: str | None = None,
    footer: str | None = None,
    add_header: bool = True,
) -> Document:
    """Aplica identidad CUN a un Document (in-place).

    - Encabezado con logo + franjas verde/lima/navy
    - Tipografía Calibri
    - Títulos recoloreados
    - Tablas con cabecera navy
    - Pie discreto
    """
    # Tipografía base
    try:
        doc.styles["Normal"].font.name = FONT
        doc.styles["Normal"].font.size = Pt(11)
        doc.styles["Normal"].font.color.rgb = RGB_GRIS
    except Exception:
        pass

    if add_header:
        # Insertar banda al inicio: construir en doc temporal y mover XML es frágil;
        # en el flujo normal convert() llama add_header ANTES del contenido.
        # Si el doc ya tiene cuerpo, anteponemos clonando elementos.
        body = doc.element.body
        # Crear banda en un doc auxiliar y copiar hijos al inicio
        tmp = Document()
        _add_header_band(tmp, subtitle=subtitle)
        # Elementos a insertar (sin sectPr)
        insert_elems = [child for child in tmp.element.body if child.tag != qn("w:sectPr")]
        ref = body[0] if len(body) else None
        for el in reversed(insert_elems):
            if ref is not None:
                ref.addprevious(el)
            else:
                body.append(el)

    style_headings_cun(doc)
    style_tables_cun(doc)

    foot = footer or "CUN · Vigilada Mineducación · Material de curso"
    _set_section_footer(doc, foot)
    return doc


def new_branded_document(
    *,
    subtitle: str | None = None,
    footer: str | None = None,
) -> Document:
    """Document vacío con encabezado de marca ya colocado (para builds MD→DOCX)."""
    doc = Document()
    try:
        doc.styles["Normal"].font.name = FONT
        doc.styles["Normal"].font.size = Pt(11)
        doc.styles["Normal"].font.color.rgb = RGB_GRIS
    except Exception:
        pass
    _add_header_band(doc, subtitle=subtitle)
    _set_section_footer(doc, footer or "CUN · Vigilada Mineducación · Material de curso")
    return doc


# Constantes exportadas para guion_md_a_docx
HEAD_BG = NAVY
ALT_BG = ALT_ROW
TITLE_COLOR = RGB_NAVY
H2_COLOR = RGB_VERDE
BODY_COLOR = RGB_GRIS
QUOTE_BG = VERDE_CLARO
