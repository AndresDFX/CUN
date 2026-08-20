# -*- coding: utf-8 -*-
"""Lo que se imprime es el `.docx`, no el `.md`. Esta comprobación lee los `.docx`.

Por grupo del martes, contra los tres artefactos impresos:
  · ficha  → tiene que llevar la pregunta NUEVA **y** la vieja (el bloque del 18/08 es a propósito)
  · hoja   → la NUEVA sí, la vieja NO (es la que se lee en voz alta)
  · índice → la NUEVA sí, la vieja NO (es lo que se repasa antes del turno)
Y que el formulario oficial siga en pie en los dos primeros.
"""
from __future__ import annotations

import os
import re
import sys

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document  # noqa: E402

RAIZ = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
FICHAS = os.path.join(RAIZ, "Especializacion", "Evaluaciones", "2026-2", "Fichas de evaluacion")
MARTES = [f"G-{n:03d}" for n in range(1, 10)]
RE_CARPETA = re.compile(r"^(\d{2}) - (G-\d{3}) - ")


def carpetas():
    out = {}
    for d in sorted(os.listdir(FICHAS)):
        m = RE_CARPETA.match(d)
        if m and os.path.isdir(os.path.join(FICHAS, d)):
            out[m.group(2)] = d
    return out


def texto_docx(p: str) -> str:
    d = Document(p)
    trozos = [x.text for x in d.paragraphs]
    for t in d.tables:
        for f in t.rows:
            trozos += [c.text for c in f.cells]
    return norm("\n".join(trozos))


def norm(s: str) -> str:
    s = (s.replace("«", "").replace("»", "").replace("“", '"').replace("”", '"')
         .replace("’", "'").replace(" ", " ").replace(" ", " ").replace(" ", " "))
    s = re.sub(r"(\d)\s+%", r"\1%", s)
    return re.sub(r"\s+", " ", s).strip()


def region(txt, ini, fin):
    m = re.search(ini + r"(.*?)(?=" + fin + r"|\Z)", txt, re.S | re.M)
    return m.group(1) if m else ""


# El extractor bueno es el de `preguntas_en_tres_sitios`: cierra cada bloque en su `»` y prefiere
# la cita al rótulo. Un `finditer` de `«…»` a secas se traga los bullets de comentario que también
# citan entre angulares («Cuidado al formularla», «díganlo con esas palabras») y da 4, 5 o 7.
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from preguntas_en_tres_sitios import preguntas as _preguntas  # noqa: E402


def citas(reg: str) -> list[str]:
    return [norm(q) for q in _preguntas(reg)]


# La huella: los primeros 12 y los últimos 12 caracteres bastan para localizar la frase en el
# docx sin depender de cómo el conversor partió los renglones intermedios.
def huella(q: str) -> tuple[str, str]:
    return q[:60], q[-60:]


def esta(q: str, doc: str) -> bool:
    a, b = huella(q)
    return a in doc and b in doc


fallos = []
for cod in MARTES:
    carp = carpetas()[cod]
    md = open(os.path.join(FICHAS, carp, "1 - Ficha de preparacion.md"), encoding="utf-8").read()
    nuevas = citas(region(md, r"Las 3 que s[íi] voy a preguntar", r"^###\s"))
    viejas = citas(region(md, r"^###[^\n]*se ley[óo] en la sala", r"^#{2,3}\s"))
    if len(nuevas) != 3 or len(viejas) != 3:
        fallos.append(f"{cod}: {len(nuevas)} nuevas / {len(viejas)} viejas en el .md")
        print(f"  ⛔ {cod}  {len(nuevas)} nuevas / {len(viejas)} viejas en el .md")
        continue

    f_doc = texto_docx(os.path.join(FICHAS, carp, "1 - Ficha de preparacion.docx"))
    h_doc = texto_docx(os.path.join(FICHAS, carp, "2 - Hoja de respuestas.docx"))
    i_doc = texto_docx(os.path.join(FICHAS, "00 - Indice y agenda de sustentaciones.docx"))

    pruebas = []
    pruebas.append(("ficha lleva las 3 nuevas", all(esta(q, f_doc) for q in nuevas)))
    pruebas.append(("ficha conserva las 3 del 18/08", all(esta(q, f_doc) for q in viejas)))
    pruebas.append(("hoja lleva las 3 nuevas", all(esta(q, h_doc) for q in nuevas)))
    pruebas.append(("hoja sin rastro de las viejas", not any(esta(q, h_doc) for q in viejas)))
    pruebas.append(("índice lleva las 3 nuevas", all(esta(q, i_doc) for q in nuevas)))
    pruebas.append(("índice sin rastro de las viejas", not any(esta(q, i_doc) for q in viejas)))
    # Los dos archivos rotulan el mismo instrumento con distinta caja: la ficha «7.1 Formulario
    # oficial del jurado» y la hoja «E · FORMULARIO OFICIAL DEL JURADO». De ahí el `re.I`; sin él
    # el control falla en los nueve grupos por igual, que es la firma de un centinela mal escrito
    # y no de una regresión.
    rotulo = re.compile(r"formulario oficial del jurado", re.I)
    pruebas.append(("formulario oficial en pie", bool(rotulo.search(f_doc))
                    and bool(rotulo.search(h_doc))))

    malas = [n for n, ok in pruebas if not ok]
    fallos += [f"{cod}: {m}" for m in malas]
    print(f"  {'⛔' if malas else '✓'} {cod}  " + (" · ".join(malas) if malas
                                                  else "los 7 controles pasan"))

print()
if fallos:
    print(f"⛔ {len(fallos)} controles fallan:")
    for f in fallos:
        print(f"   {f}")
else:
    print("✓ los .docx impresos llevan las preguntas nuevas; la hoja y el índice sin la vieja,")
    print("  y la ficha conserva las tres del 18/08 para cotejar la hoja escrita a mano.")
raise SystemExit(1 if fallos else 0)
