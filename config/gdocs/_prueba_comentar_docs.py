# -*- coding: utf-8 -*-
r"""
Prueba `comentar_docs.py` y `syllabus_curso.py` sin cuenta, sin red y sin tocar ningún documento.

Por qué existe: la mitad peligrosa de este proceso escribe **en el documento de un estudiante**, y un
comentario publicado le llega por correo al instante y no se deshace sin dejar rastro. La red de
seguridad son las tres puertas antes de publicar —`validar`, `revisar_lenguaje` y `ensayar`— y esas
puertas son código que también se puede romper. Aquí se prueban con documentos de mentira construidos
a mano, incluida la trampa que ya nos costó un comentario mal puesto: la **tabla de contenido**.

Cubre lo que de verdad se puede romper:
  1. `doc_id` — enlace, enlace con parámetros, id pelado, basura
  2. `_sanear` — el `.docx.docx`, los caracteres que Windows no admite, el nombre vacío
  3. `revisar_lenguaje` — los ocho vicios de oficio, y que la prosa natural pasa
  4. la tabla de contenido — `parrafos_docx` no la ve y Ctrl+F sí: contar con la vista equivocada
     deja pasar una cita duplicada
  5. `validar` — cita corta · no literal · que cruza párrafo · duplicada por el índice · criterio
     inventado · comentario vacío · plan sin comentarios · el camino feliz
  6. `validar` en tabla — el cronograma y el presupuesto van en tabla y también se comentan
  7. `leer_plan` — falta `curso`, `aca` o `comentarios`
  8. `criterios_aca` — las seis guías siguen dando sus 41 criterios
  9. `syllabus_curso` — los dos formatos, las unidades que van de dos en dos, y el tercer formato de TG2

    python config/gdocs/_prueba_comentar_docs.py
"""
from __future__ import annotations

import json
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

for _flujo in (sys.stdout, sys.stderr):
    if hasattr(_flujo, "reconfigure"):
        _flujo.reconfigure(encoding="utf-8", errors="replace")

sys.path.insert(0, str(Path(__file__).resolve().parent))
import comentar_docs as CD  # noqa: E402
import criterios_aca as CA  # noqa: E402
import plan_comentarios as PC  # noqa: E402
import syllabus_curso as SY  # noqa: E402

CRITERIOS_TG2 = CA.criterios("tg2", "acafinal")[1]


# ── documentos de mentira ────────────────────────────────────────────────────

def docx_de_prueba(destino: Path, parrafos: list[str], *, indice: list[str] | None = None,
                   tabla: list[list[str]] | None = None) -> Path:
    """Un `.docx` con lo que se le pida, y con el índice envuelto en `w:sdt` **como lo hace Google**.

    Es el detalle que importa: python-docx no puede crear un `w:sdt`, así que se escribe el documento
    con python-docx y después se le inyecta el bloque en el XML, que es exactamente la forma en que
    llega un documento exportado de Google Docs con tabla de contenido.
    """
    from docx import Document

    d = Document()
    for t in parrafos:
        d.add_paragraph(t)
    if tabla:
        t = d.add_table(rows=len(tabla), cols=len(tabla[0]))
        for fila, celdas in zip(t.rows, tabla):
            for celda, texto in zip(fila.cells, celdas):
                celda.text = texto
    d.save(str(destino))

    if not indice:
        return destino

    sdt = "<w:sdt><w:sdtContent>" + "".join(
        f"<w:p><w:r><w:t>{t}</w:t></w:r></w:p>" for t in indice) + "</w:sdtContent></w:sdt>"
    with zipfile.ZipFile(destino) as z:
        partes = {n: z.read(n) for n in z.namelist()}
    xml = partes["word/document.xml"].decode("utf-8")
    xml = xml.replace("<w:body>", "<w:body>" + sdt, 1)
    partes["word/document.xml"] = xml.encode("utf-8")
    with zipfile.ZipFile(destino, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in partes.items():
            z.writestr(n, b)
    return destino


def plan_de(cita: str, texto: str = "Aquí falta el paso que convierte estos dos números en la "
                                    "conclusión que sacas enseguida.",
            criterio: str | None = None) -> dict:
    return {"curso": "tg2", "aca": "acafinal",
            "comentarios": [{"criterio": criterio or CRITERIOS_TG2[0], "cita": cita, "texto": texto}]}


def main() -> int:
    fallos: list[str] = []

    def check(nombre: str, cond: bool, detalle: str = "") -> None:
        print(("  OK   " if cond else "  FALLA ") + nombre + ("" if cond else f"  → {detalle}"))
        if not cond:
            fallos.append(nombre)

    tmp = Path(tempfile.mkdtemp(prefix="prueba-comentar-"))
    try:
        # ── 1. doc_id ────────────────────────────────────────────────────────
        ID = "1KaQdgXIAtqzJ7pCLgZrielxFJYr5uzCe"
        check("1. doc_id acepta el enlace compartido",
              CD.doc_id(f"https://docs.google.com/document/d/{ID}/edit") == ID)
        check("1b. …con parámetros detrás",
              CD.doc_id(f"https://docs.google.com/document/d/{ID}/edit?usp=sharing&tab=t.0") == ID)
        check("1c. …y el id pelado", CD.doc_id(f"  {ID}  ") == ID)
        try:
            CD.doc_id("el documento de Harold")
            check("1d. basura aborta en vez de seguir", False, "no abortó")
        except SystemExit:
            check("1d. basura aborta en vez de seguir", True)

        # ── 2. _sanear ───────────────────────────────────────────────────────
        check("2. _sanear no deja .docx.docx",
              CD._sanear("HAROLD HURTADO Plantilla Anteproyectos.docx")
              == "HAROLD HURTADO Plantilla Anteproyectos",
              CD._sanear("HAROLD HURTADO Plantilla Anteproyectos.docx"))
        check("2b. quita los caracteres que Windows no admite",
              # / : " " ? → cinco guiones, incluido el del final
              CD._sanear('TG2/TG3: informe "final"?') == "TG2-TG3- informe -final--",
              CD._sanear('TG2/TG3: informe "final"?'))
        check("2c. un título vacío no deja un nombre vacío", CD._sanear("   ") == "documento")
        check("2d. corta a 110 caracteres", len(CD._sanear("x" * 300)) == 110)

        # ── 3. revisar_lenguaje ──────────────────────────────────────────────
        oficio = [
            "Criterio «Documento integrado»: no se cumple.",
            "No cumple con lo pedido en la guía.",
            "Según el criterio de evaluación, falta la metodología.",
            "De acuerdo con la rúbrica, esto está incompleto.",
            "Se evidencia ausencia de antecedentes actualizados.",
            "El estudiante no desarrolló el marco teórico.",
            "Párrafo 12: falta la cita.",
            "El desarrollo es satisfactorio pero mejorable.",
        ]
        for t in oficio:
            check(f"3. «{t[:38]}…» se rechaza", bool(CD.revisar_lenguaje(t)))
        natural = ("Aquí dices 32 y en Población y muestra dices 34. Los porcentajes de las figuras "
                   "solo cuadran con 34: el 41,2% son 14 de 34. Deja el número con el que calculaste.")
        check("3b. la prosa natural pasa", not CD.revisar_lenguaje(natural),
              str(CD.revisar_lenguaje(natural)))

        # ── 4. la tabla de contenido ─────────────────────────────────────────
        CUERPO = ("Los hallazgos obtenidos durante el desarrollo de la investigación permiten "
                  "identificar tres asuntos.")
        ENCABEZADO = "Interpretación Técnica de Hallazgos"
        doc = docx_de_prueba(tmp / "con-indice.docx",
                             ["Introducción del trabajo con su texto de relleno.",
                              ENCABEZADO, CUERPO],
                             indice=[f"PAGEREF _heading=h.abc \\h {ENCABEZADO}", "Conclusiones"])
        ps = PC.parrafos_docx(doc)
        visibles = CD.texto_como_lo_busca_ctrl_f(doc)
        check("4. parrafos_docx NO ve la tabla de contenido",
              sum(1 for p in ps if ENCABEZADO in p["texto"]) == 1,
              f"{sum(1 for p in ps if ENCABEZADO in p['texto'])} apariciones")
        check("4b. texto_como_lo_busca_ctrl_f SÍ la ve",
              sum(1 for t in visibles if ENCABEZADO in t) == 2,
              f"{sum(1 for t in visibles if ENCABEZADO in t)} apariciones")

        # ── 5. validar ───────────────────────────────────────────────────────
        pr, av = CD.validar(plan_de(ENCABEZADO), ps, visibles, CRITERIOS_TG2)
        check("5. una cita que también está en el índice se rechaza",
              any("2 veces" in x and "tabla de contenido" in x for x in pr), " | ".join(pr)[:220])
        pr, av = CD.validar(plan_de(CUERPO[:70]), ps, visibles, CRITERIOS_TG2)
        check("5b. la misma frase, citada del cuerpo, pasa", not pr, " | ".join(pr)[:220])
        pr, _ = CD.validar(plan_de("tres asuntos"), ps, visibles, CRITERIOS_TG2)
        check("5c. una cita corta se rechaza por corta",
              any(f"mínimo {CD.CITA_MIN}" in x for x in pr), " | ".join(pr)[:200])
        pr, _ = CD.validar(plan_de("Los  hallazgos obtenidos durante el desarrollo"),
                           ps, visibles, CRITERIOS_TG2)
        check("5d. un espacio doble se rechaza: Ctrl+F busca literal",
              any("NO letra por letra" in x for x in pr), " | ".join(pr)[:220])
        pr, _ = CD.validar(plan_de(f"{ENCABEZADO} Los hallazgos obtenidos durante el desarrollo"),
                           ps, visibles, CRITERIOS_TG2)
        check("5e. una cita que cruza un salto de párrafo se rechaza y lo explica",
              any("cruza un salto de párrafo" in x for x in pr), " | ".join(pr)[:220])
        pr, _ = CD.validar(plan_de(CUERPO[:70], criterio="Redacción impecable"),
                           ps, visibles, CRITERIOS_TG2)
        check("5f. un criterio que no está en la guía se rechaza, y se enseña la guía",
              any("no está en la guía" in x for x in pr) and any(CRITERIOS_TG2[0] in x for x in pr),
              " | ".join(pr)[:200])
        pr, _ = CD.validar(plan_de(CUERPO[:70], texto=""), ps, visibles, CRITERIOS_TG2)
        check("5g. un comentario vacío se rechaza",
              any("sin `texto`" in x for x in pr), " | ".join(pr)[:200])
        pr, _ = CD.validar({"curso": "tg2", "aca": "acafinal", "comentarios": []},
                           ps, visibles, CRITERIOS_TG2)
        check("5h. un plan sin comentarios se rechaza",
              any("no tiene comentarios" in x for x in pr), " | ".join(pr)[:200])
        plan = plan_de(CUERPO[:70])
        pr, _ = CD.validar(plan, ps, visibles, CRITERIOS_TG2)
        check("5i. el camino feliz deja anotado en qué párrafo cae",
              plan["comentarios"][0].get("_parrafo") is not None,
              json.dumps(plan["comentarios"][0], ensure_ascii=False)[:200])

        # ── 6. citas dentro de una tabla ─────────────────────────────────────
        EN_TABLA = "Aplicación del instrumento a la muestra seleccionada"
        doc2 = docx_de_prueba(tmp / "con-tabla.docx", ["Cronograma de actividades del proyecto."],
                              tabla=[["Actividad", "Mes"], [EN_TABLA, "Marzo"]])
        ps2 = PC.parrafos_docx(doc2)
        vis2 = CD.texto_como_lo_busca_ctrl_f(doc2)
        pr, _ = CD.validar(plan_de(EN_TABLA), ps2, vis2, CRITERIOS_TG2)
        check("6. una cita del cronograma, que va en tabla, se puede comentar", not pr,
              " | ".join(pr)[:220])

        # ── 7. leer_plan ─────────────────────────────────────────────────────
        for falta in ("curso", "aca", "comentarios"):
            p = plan_de(CUERPO[:70])
            p.pop(falta)
            ruta = tmp / f"sin-{falta}.json"
            ruta.write_text(json.dumps(p, ensure_ascii=False), encoding="utf-8")
            try:
                CD.leer_plan(ruta)
                check(f"7. un plan sin «{falta}» aborta", False, "no abortó")
            except SystemExit as e:
                check(f"7. un plan sin «{falta}» aborta", falta in str(e), str(e))

        # ── 8. criterios_aca ─────────────────────────────────────────────────
        todo = CA.todos()
        cuenta = sum(len(v["criterios"]) for c in todo.values() for v in c.values())
        guias = sum(len(c) for c in todo.values())
        check("8. las seis guías de ACA siguen dando 41 criterios",
              (guias, cuenta) == (6, 41), f"{guias} guías, {cuenta} criterios")
        check("8b. ninguna guía se quedó sin criterios",
              all(v["criterios"] for c in todo.values() for v in c.values()),
              str([(k, s) for k, c in todo.items() for s, v in c.items() if not v["criterios"]]))

        # ── 9. syllabus_curso ────────────────────────────────────────────────
        # TG2 SÍ tiene Syllabus desde el 22/08/2026. Y es el caso raro: llegó **sin el código SIAC
        # en el nombre** (su casilla «CÓDIGO SÍAC» está vacía), así que se localiza por el respaldo
        # estructural; y trae el armazón SIAC de pregrado con la tabla de unidades **didácticas**
        # del formato de especialización. Las dos cosas se comprueban aquí porque las dos costaron.
        f = SY.ficha("tg2")
        check("9. TG2 ya tiene Syllabus, y se encuentra pese a no llevar el código en el nombre",
              f["archivo"] is not None and "94453" not in Path(f["archivo"]).name,
              json.dumps(f, ensure_ascii=False)[:200])
        # El número era 6 hasta que el lector aprendió a recorrer el armazón SIAC completo (commit
        # «TG2 lee sus 12 unidades»). La prueba se quedó con el 6 y llevaba fallando desde entonces:
        # un arnés que falla siempre por lo mismo deja de avisar de lo que sí importa.
        check("9b. TG2: 12 unidades didácticas dentro del armazón SIAC, y el artículo de reflexión",
              len(f["unidades"]) == 12
              and [u["n"] for u in f["unidades"]] == list(range(1, 13))
              and "artículo de reflexión" in (f["competencia"] or "")
              and "CORTE ÚNICO" in (f["evaluacion"] or "").upper(),
              f"{len(f['unidades'])} unidades · {(f['competencia'] or '')[:60]}")
        f = SY.ficha("tg3")
        check("9c. TG3: formato SIAC, 14 unidades, competencia y resultado",
              (f["formato"], len(f["unidades"])) == ("siac", 14) and f["competencia"] and f["resultado"],
              f"{f['formato']} · {len(f['unidades'])} unidades")
        f = SY.ficha("investigacion")
        combinadas = [u for u in f["unidades"] if u.get("tambien")]
        numeros = sorted({u["n"] for u in f["unidades"]} |
                         {x for u in f["unidades"] for x in u.get("tambien", [])})
        check("9d. Investigación: las unidades que van de dos en dos no se pierden",
              numeros == [1, 2, 3, 4, 5, 6, 7, 8, 10, 11, 12] and len(combinadas) == 1, str(numeros))
        check("9e. …y se avisa de que el Syllabus se salta la U9",
              any("salta la unidad 9" in x for x in f["advertencias"]), str(f["advertencias"]))
        f = SY.ficha("proyecto1")
        check("9f. Proyecto I: el otro formato, 7 unidades didácticas",
              (f["formato"], len(f["unidades"])) == ("especializacion", 7) and f["competencia"],
              f"{f['formato']} · {len(f['unidades'])} unidades")
        check("9g. …y se avisa de que su numeración ACA 1/2/3 no es la del aula",
              any("NO es la del aula" in x for x in f["advertencias"]), str(f["advertencias"]))
        f = SY.ficha("creatividad")
        check("9h. Creatividad: 8 unidades y 3 elementos de competencia",
              (len(f["unidades"]), len(f["elementos"])) == (8, 3),
              f"{len(f['unidades'])} unidades, {len(f['elementos'])} elementos")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    print(f"\n{len(fallos)} fallos")
    return 1 if fallos else 0


if __name__ == "__main__":
    raise SystemExit(main())
