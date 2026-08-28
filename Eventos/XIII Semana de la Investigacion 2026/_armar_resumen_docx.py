# -*- coding: utf-8 -*-
"""Arma el resumen de memorias de la XIII Semana de la Investigación en UNA sola página.

El convertidor genérico (`config/slides/guion_md_a_docx.py`) sirve para guiones y fichas, donde
sobra el espacio; aquí lo derramaba a una segunda hoja para dos líneas. Este armador es específico
del formato de memorias: el comité compila los resúmenes en un solo documento, y un resumen de 300
palabras que ocupa dos hojas se ve como un descuido.

La fuente de verdad sigue siendo el `.md` hermano: este script lo lee, no duplica el texto.

Uso:  python _armar_resumen_docx.py
Comprobar:  soffice --headless --convert-to pdf  → tiene que dar 1 página.
"""
from __future__ import annotations

import io
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

AQUI = os.path.dirname(os.path.abspath(__file__))
BASE = "Castaño Espinosa_Julian Andrés"

FUENTE = "Calibri"
VERDE = RGBColor(0x00, 0x74, 0x33)   # verde marca CUN
AZUL = RGBColor(0x0C, 0x23, 0x40)    # azul marino CUN
GRIS = RGBColor(0x59, 0x59, 0x59)    # solo para la cabecera de plantilla, que va discreta
# El cuerpo va casi negro a propósito: el gris de marca (#595959) se lee lavado impreso, y esto
# entra a unas memorias que alguien va a leer en papel.
TINTA = RGBColor(0x1A, 0x1A, 0x1A)


def _leer_md(ruta: str) -> dict:
    """Saca del markdown la cabecera, los 5 campos con etiqueta, el resumen y las palabras clave."""
    txt = io.open(ruta, encoding="utf-8").read()
    cabecera = re.search(r"^#\s+(.+)$", txt, re.M).group(1)
    campos = re.findall(r"^\*\*(.+?):\*\*\s*(.+)$", txt, re.M)
    resumen = txt.split("## RESUMEN")[1].split("## PALABRAS CLAVE")[0]
    parrafos = [p.strip() for p in resumen.strip().split("\n\n") if p.strip()]
    claves = txt.split("## PALABRAS CLAVE")[1].strip()
    return {"cabecera": cabecera, "campos": campos, "parrafos": parrafos, "claves": claves}


def _par(doc, texto, *, size, bold=False, color=GRIS, align=None,
         antes=0, despues=0, interlineado=1.0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(antes)
    pf.space_after = Pt(despues)
    pf.line_spacing = interlineado
    if align is not None:
        p.alignment = align
    r = p.add_run(texto)
    r.font.name = FUENTE
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    return p


def _par_etiquetado(doc, etiqueta, valor):
    """«Autor(es): Julian…» — la etiqueta en negrita azul, el valor en gris."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    for texto, bold, color in ((f"{etiqueta}: ", True, AZUL), (valor, False, TINTA)):
        r = p.add_run(texto)
        r.font.name = FUENTE
        r.font.size = Pt(10.5)
        r.font.bold = bold
        r.font.color.rgb = color
    return p


def armar() -> str:
    md = os.path.join(AQUI, BASE + ".md")
    out = os.path.join(AQUI, BASE + ".docx")
    d = _leer_md(md)

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.8)
        s.left_margin = s.right_margin = Cm(2.2)

    # La cabecera institucional es de la plantilla, no del trabajo: va discreta arriba.
    _par(doc, d["cabecera"], size=8.5, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, despues=8)

    titulo = dict(d["campos"])["Título del trabajo"]
    _par(doc, titulo, size=13, bold=True, color=VERDE,
         align=WD_ALIGN_PARAGRAPH.CENTER, despues=8, interlineado=1.05)

    for etiqueta, valor in d["campos"]:
        if etiqueta == "Título del trabajo":
            continue
        _par_etiquetado(doc, etiqueta, valor)

    _par(doc, "RESUMEN", size=11.5, bold=True, color=VERDE, antes=10, despues=4)
    for parrafo in d["parrafos"]:
        _par(doc, parrafo, size=11, color=TINTA, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             despues=6, interlineado=1.08)

    _par(doc, "PALABRAS CLAVE", size=11.5, bold=True, color=VERDE, antes=8, despues=4)
    _par(doc, d["claves"], size=11, color=TINTA, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         interlineado=1.08)

    doc.save(out)
    return out


if __name__ == "__main__":
    print("OK", armar())
