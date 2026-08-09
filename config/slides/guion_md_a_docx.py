# -*- coding: utf-8 -*-
"""
Convierte Markdown a .docx bien formateado.

Nota CUN: los guiones docentes en `Guiones/` son **solo .md** (no usar este
script para regenerar docx de guiones). Sí aplica para material de estudiantes
en `Clases/` (LEEME, fichas, ACAs) vía `sync_clases_estudiantes.py` /
`build_acas_estudiantes.py`.

Por defecto aplica **marca CUN** (`cun_docx_brand`: logo, #007433 / #91DC00 /
#0C2340, tipografía Calibri). Desactivar con ``brand=False``.

Uso:  python guion_md_a_docx.py "ruta/al/guion.md"  [--out "ruta.docx"] [--no-brand]
"""
from __future__ import annotations

import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from cun_docx_brand import (
    ALT_BG,
    BODY_COLOR,
    FONT,
    H2_COLOR,
    HEAD_BG,
    QUOTE_BG,
    TITLE_COLOR,
    apply_cun_brand_docx,
    new_branded_document,
)

DARK = TITLE_COLOR
GRAY = BODY_COLOR
ORANGE = RGBColor(0xC2, 0x3E, 0x12)
CODE_BG = "F2F2F3"
MONO = "Consolas"

CAPTURAS = ""


def _shade(el, fill):
    pPr = el.get_or_add_pPr() if el.tag.endswith("}p") else el.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def shade_par(p, fill):
    _shade(p._p, fill)


def shade_cell(c, fill):
    tcPr = c._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_inline(p, text, base_size=11, base_color=GRAY, base_bold=False):
    """Parseo inline: **negrita** y `código`."""
    for part in re.split(r"(\*\*.*?\*\*|`[^`]+`)", text):
        if not part:
            continue
        r = p.add_run()
        if part.startswith("**") and part.endswith("**"):
            r.text = part[2:-2]
            r.bold = True
            r.font.size = Pt(base_size)
            r.font.color.rgb = base_color
            r.font.name = FONT
        elif part.startswith("`") and part.endswith("`"):
            r.text = part[1:-1]
            r.font.name = MONO
            r.font.size = Pt(base_size - 0.5)
            r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        else:
            r.text = part
            r.bold = base_bold
            r.font.size = Pt(base_size)
            r.font.color.rgb = base_color
            r.font.name = FONT


def find_capture(caption):
    c = caption.lower()
    base = CAPTURAS
    if not base or not os.path.isdir(base):
        return None

    def pick(sub):
        for f in sorted(os.listdir(base)):
            if sub in f.lower() and f.lower().endswith(".png"):
                return os.path.join(base, f)
        return None

    if "subnet" in c or "subred" in c:
        return pick("subnetting")
    if "jslinux" in c or "webminal" in c or ("terminal" in c and "linux" in c) or "distrosea" in c or "copy.sh" in c:
        return pick("jslinux")
    if re.search(
        r"(inicia\w*\s+sesi|inicio de sesi|iniciar sesi|\blogin\b|autentic|"
        r"(acces\w*|acceder|entrar|ingres\w*|iniciar)\s+(a\s+)?examlab|pantalla de acceso)",
        c,
    ):
        return pick("login-form") or pick("login")
    return None


def _capture_token(text):
    m = re.search(r"\[\[\s*captura:\s*([^\]]+?)\s*\]\]", text, re.I)
    if not m:
        return text, None
    fn = m.group(1).strip()
    if not fn.lower().endswith(".png"):
        fn += ".png"
    return re.sub(r"\[\[\s*captura:[^\]]+\]\]", "", text, flags=re.I).strip(), fn


def screenshot_box(doc, caption, explicit=None):
    p = doc.add_paragraph()
    shade_par(p, "FFF4EC")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("📸  " + caption)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = ORANGE
    img = None
    if explicit:
        cand = os.path.join(CAPTURAS, explicit)
        if os.path.isfile(cand):
            img = cand
    if img is None:
        img = find_capture(caption)
    if img:
        try:
            doc.add_picture(img, width=Inches(5.6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
    else:
        ph = doc.add_paragraph()
        shade_par(ph, CODE_BG)
        rr = ph.add_run("[  Inserta aquí la captura de pantalla  ]")
        rr.italic = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = GRAY


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    shade_par(p, CODE_BG)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for i, ln in enumerate(lines):
        if i:
            p.add_run().add_break()
        r = p.add_run(ln)
        r.font.name = MONO
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)


def add_heading(doc, level, text, brand=True):
    sizes = {1: 20, 2: 16, 3: 13.5, 4: 12, 5: 11}
    color = TITLE_COLOR if level <= 1 else (H2_COLOR if level == 2 and brand else DARK)
    if brand and level >= 3:
        color = TITLE_COLOR
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level <= 2 else 7)
    p.paragraph_format.space_after = Pt(3)
    add_inline(p, text, base_size=sizes.get(level, 12), base_color=color, base_bold=True)
    for r in p.runs:
        r.bold = True


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def parse_table_row(line):
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def convert(
    md_path,
    out_path,
    *,
    brand: bool = True,
    footer: str | None = None,
    subtitle: str | None = None,
):
    """MD → DOCX. ``brand=True`` (default): identidad CUN en todo el documento."""
    global CAPTURAS
    CAPTURAS = os.path.join(os.path.dirname(os.path.abspath(md_path)), "Capturas")
    with open(md_path, encoding="utf-8") as f:
        raw = f.read().split("\n")

    if brand:
        doc = new_branded_document(subtitle=subtitle, footer=footer)
    else:
        doc = Document()
        doc.styles["Normal"].font.name = FONT
        doc.styles["Normal"].font.size = Pt(11)

    head_bg = HEAD_BG if brand else "3A3A3C"
    alt_bg = ALT_BG if brand else "ECECEE"
    title_done = False
    i = 0
    while i < len(raw):
        line = raw[i]
        s = line.strip()
        if s.startswith("```"):
            block = []
            i += 1
            while i < len(raw) and not raw[i].strip().startswith("```"):
                block.append(raw[i])
                i += 1
            add_code_block(doc, block)
            i += 1
            continue
        if s.startswith("|") and i + 1 < len(raw) and is_table_sep(raw[i + 1]):
            header = parse_table_row(raw[i])
            i += 2
            rows = []
            while i < len(raw) and raw[i].strip().startswith("|"):
                rows.append(parse_table_row(raw[i]))
                i += 1
            ncol = len(header)
            t = doc.add_table(rows=1 + len(rows), cols=ncol)
            t.style = "Table Grid"
            for j, h in enumerate(header[:ncol]):
                c = t.cell(0, j)
                c.paragraphs[0].text = ""
                shade_cell(c, head_bg)
                add_inline(
                    c.paragraphs[0],
                    h,
                    base_size=10.5,
                    base_color=RGBColor(0xFF, 0xFF, 0xFF),
                    base_bold=True,
                )
                for r in c.paragraphs[0].runs:
                    r.bold = True
            for ri, row in enumerate(rows):
                for j in range(ncol):
                    c = t.cell(ri + 1, j)
                    c.paragraphs[0].text = ""
                    if ri % 2 == 1:
                        shade_cell(c, alt_bg)
                    add_inline(c.paragraphs[0], row[j] if j < len(row) else "", base_size=10)
            doc.add_paragraph()
            continue
        if not s:
            i += 1
            continue
        if re.match(r"^(---+|\*\*\*+|___+)$", s):
            i += 1
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            hashes, text = m.group(1), m.group(2)
            lvl = len(hashes)
            if not title_done:
                p = doc.add_paragraph()
                add_inline(p, text, base_size=22, base_color=TITLE_COLOR if brand else DARK, base_bold=True)
                for r in p.runs:
                    r.bold = True
                title_done = True
            else:
                add_heading(doc, max(2, lvl - 1), text, brand=brand)
            i += 1
            continue
        if s.startswith(">"):
            content = s.lstrip(">").strip()
            if "📸" in content or "🖼" in content or re.search(r"\[\[\s*captura:", content, re.I):
                content, explicit = _capture_token(content)
                cap = (
                    re.sub(r"^\s*(📸|🖼️?)\s*", "", content)
                    .replace("**", "")
                    .replace("Pantallazo —", "")
                    .replace("Imagen —", "")
                    .strip()
                )
                cap = re.sub(r"^\W+", "", cap)
                screenshot_box(doc, cap or "Imagen", explicit)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                shade_par(p, QUOTE_BG if brand else "F7F7F8")
                add_inline(p, content, base_size=10.5, base_color=GRAY)
                for r in p.runs:
                    r.italic = True
            i += 1
            continue
        if (
            s.startswith("📸")
            or s.startswith("**📸")
            or "📸 **Pantallazo" in s
            or s.startswith("🖼")
            or re.search(r"\[\[\s*captura:", s, re.I)
        ):
            s2, explicit = _capture_token(s)
            cap = (
                re.sub(r"^\W*(📸|🖼️?)\W*", "", s2)
                .replace("**", "")
                .replace("Pantallazo —", "")
                .replace("Imagen —", "")
                .strip()
            )
            screenshot_box(doc, cap or "Imagen", explicit)
            i += 1
            continue
        mli = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if mli:
            indent = len(mli.group(1))
            ordered = bool(re.match(r"\d+\.", mli.group(2)))
            style = "List Number" if ordered else "List Bullet"
            if indent >= 2 and not ordered:
                style = "List Bullet 2"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25 + (0.25 if indent >= 2 else 0))
            add_inline(p, mli.group(3), base_size=11)
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, s, base_size=11)
        i += 1

    if brand:
        # Refuerza tablas/títulos sin duplicar header
        apply_cun_brand_docx(doc, subtitle=subtitle, footer=footer, add_header=False)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    src = sys.argv[1]
    out = None
    brand = "--no-brand" not in sys.argv
    if "--out" in sys.argv:
        out = sys.argv[sys.argv.index("--out") + 1]
    else:
        out = os.path.splitext(src)[0] + ".docx"
    convert(src, out, brand=brand)
    print("OK", out)
