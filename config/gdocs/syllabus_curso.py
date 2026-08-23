# -*- coding: utf-8 -*-
"""
El Syllabus de la asignatura, leído del `.docx` oficial que está en la raíz del curso.

POR QUÉ HACE FALTA, SI YA HAY CRITERIOS
---------------------------------------
`criterios_aca.py` trae el checklist de la guía del ACA: **qué se le pidió** al estudiante en esa
entrega. El Syllabus trae otra cosa: **qué enseña la asignatura** — su competencia, sus elementos de
competencia, las unidades de conocimiento y el resultado de aprendizaje esperado. Al comentar, lo
primero dice si la entrega está completa y lo segundo si el documento hace lo que el curso pretende
que la persona sepa hacer. Un comentario que solo mira el checklist se convierte en un inventario;
uno que además mira el Syllabus puede decir por qué eso importa más allá de la nota.

CÓMO SE LOCALIZA EL ARCHIVO
---------------------------
Por el **código SIAC**, que ya vive en `config/cursos/sesiones_cun.py` (fuente única: ahí están el
código, el temario por sesión y la nota de qué falta). El Syllabus lleva el código en el nombre
(`…EI005_PRES.docx`, `…94532_PRES_VIR.docx`, `…ESP329.docx`), así que se busca por eso y no por una
tabla de nombres que se desactualiza sola.

**TG2 no tiene Syllabus** y eso no es un fallo del programa: la CUN nunca lo entregó, y el
repositorio lo dice abiertamente en el Manual del Docente y en `sesiones_cun.py` (`nota_syllabus`).
Aquí se devuelve `archivo=None` con la nota, y quien redacta los comentarios lo sabe en vez de
creerse un Syllabus inventado.

HAY DOS FORMATOS, NO UNO
------------------------
1. **SIAC de pregrado** (EI004, EI005, 94532): *todo* tablas —19 tablas y un único párrafo con
   texto—. Las secciones se reconocen por su título («2. DESCRIPCIÓN…», «UNIDADES DE CONOCIMIENTO»)
   en cualquier fila, no por el número de tabla: hay secciones que comparten tabla (el propósito y la
   competencia van juntos). Y **una fila puede traer dos unidades**: en Investigación el número dice
   «3 4» y «11 12», porque esas unidades se dictan en una sola sesión.
2. **Sílabo de especialización** (ESP329): al contrario, casi todo párrafos con encabezados, y las
   unidades en una tabla «Unidad didáctica · ¿Qué debe enseñar el docente? · Producto esperado».

Se detecta la forma y se lee cada una como es. Si aparece un tercer formato, `ficha()` devuelve las
claves vacías y `advertencias` lo dice; no adivina.
"""
from __future__ import annotations

import glob
import re
import sys
from pathlib import Path

for _flujo in (sys.stdout, sys.stderr):
    if hasattr(_flujo, "reconfigure"):
        _flujo.reconfigure(encoding="utf-8", errors="replace")

REPO = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO / "config" / "cursos"))

import sesiones_cun as SC  # noqa: E402  (después de tocar sys.path)

# ── formato 1: SIAC de pregrado ────────────────────────────────────────────────
# Las secciones que sirven para comentar. El resto —control de cambios, perfil docente,
# aprobaciones— es trámite y no se lee.
SECCIONES_SIAC: list[tuple[str, str]] = [
    ("descripcion", r"DESCRIPCI[ÓO]N DE LA UNIDAD CURRICULAR"),
    ("proposito", r"PROP[ÓO]SITO DE FORMACI[ÓO]N"),
    ("competencia", r"COMPETENCIA DE LA UNIDAD CURRICULAR"),
    ("elementos", r"ELEMENTOS DE COMPETENCIA"),
    ("_tramite", r"PRERREQUISITOS Y CORREQUISITOS"),
    ("unidades", r"UNIDADES DE CONOCIMIENTO"),
    ("_tramite", r"\bESTRATEGIAS\b"),
    ("bibliografia", r"RECURSOS BIBLIOGR[ÁA]FICOS"),
    ("resultado", r"RESULTADO DE APRENDIZAJE"),
    ("mecanismos", r"MECANISMOS DE EVALUACI[ÓO]N"),
    ("evaluacion", r"SISTEMA DE EVALUACI[ÓO]N"),
    ("_tramite", r"PERFIL DOCENTE|HIST[ÓO]RICO DE CAMBIOS|CONTROL DE (APROBACI|CAMBIOS)"),
]

# ── formato 2: sílabo de especialización ───────────────────────────────────────
SECCIONES_ESP: list[tuple[str, str]] = [
    ("_tramite", r"^\s*INFORMACI[ÓO]N GENERAL"),
    ("descripcion", r"^\s*JUSTIFICACI[ÓO]N"),
    ("_tramite", r"^\s*ARTICULACI[ÓO]N CON PROYECTO"),
    ("competencia", r"MACROCOMPETENCIA"),
    ("_tramite", r"^\s*ESTRUCTURA TEM[ÁA]TICA"),
    ("estrategia", r"^\s*ESTRATEGIA DID[ÁA]CTICA"),
    ("mecanismos", r"MECANISMOS Y ESTRATEGIAS DE EVALUACI[ÓO]N"),
    ("_tramite", r"^\s*(Alternativas metodol|Medios y ayudas)"),
    ("bibliografia", r"^\s*BIBLIOGRAF[ÍI]A"),
    ("_tramite", r"^\s*\d*\.?\s*OBSERVACIONES"),
]

TEXTO = ("descripcion", "proposito", "competencia", "resultado", "mecanismos", "estrategia")


def archivo(curso: str) -> Path | None:
    """El `.docx` del Syllabus del curso, o `None` si no está en la carpeta.

    Primero por **código SIAC**, que es preciso. Pero el de TG2 llegó sin el código en el nombre
    (`TRABJO DE GRADO II INGENIERIA DE SISTEMAS.docx`, y su casilla «CÓDIGO SÍAC» está vacía), así
    que hay un respaldo: se abre cada `.docx` y se acepta el que traiga el armazón del formato SIAC.
    No vale «el único `.docx` de la carpeta» —la de TG3 guarda además el correo de bienvenida—, y
    por eso el respaldo mira **dentro**: «UNIDAD CURRICULAR» está en la primera tabla de todos los
    Syllabus y en ningún otro documento del curso.
    """
    c = SC.COURSES[curso]
    codigo = str(c.get("codigo") or "")
    candidatos = [Path(x) for x in sorted(glob.glob(str(Path(c["folder"]) / "*.docx")))]
    for p in candidatos:
        if codigo and codigo in p.name:
            return p
    for p in candidatos:
        if p.name.startswith("~$"):          # bloqueo de Word abierto
            continue
        try:
            from docx import Document          # local: el módulo se importa sin python-docx puesto
            doc = Document(str(p))
        except Exception:
            continue
        for tabla in doc.tables[:3]:
            if "UNIDAD CURRICULAR" in " ".join(
                    celda.text.upper() for fila in tabla.rows for celda in fila.cells):
                return p
    return None


def _filas(tabla) -> list[list[str]]:
    """Las filas de la tabla sin las repeticiones de las celdas combinadas.

    En el formato SIAC casi toda cabecera es una celda combinada a lo ancho, y python-docx la
    devuelve repetida en cada columna: sin esto, el título sale seis veces.
    """
    fuera = []
    for r in tabla.rows:
        celdas, visto = [], None
        for c in r.cells:
            t = " ".join(c.text.split())
            if t and t == visto:
                continue
            visto = t
            celdas.append(t)
        fuera.append(celdas)
    return fuera


def _seccion_de(texto: str, secciones) -> str | None:
    t = " ".join(texto.split()).upper()
    if len(t) > 160:                    # una cabecera es corta; esto ya es contenido
        return None
    for clave, patron in secciones:
        if re.search(patron, t, re.I):
            return clave
    return None


def _es_rotulo(fila: list[str]) -> bool:
    """¿Es la fila de rótulos («N° | TEMÁTICA | SUBTEMÁTICA»)? No aporta contenido."""
    t = " ".join(" ".join(fila).split()).upper().strip(" |")
    return t in {"N°", "Nº", "N"} or bool(re.fullmatch(
        r"(N[°º]?\s*)?(TEM[ÁA]TICA|SUBTEM[ÁA]TICA|RELACI[ÓO]N DE ELEMENTOS.*|REFERENCIAS.*|"
        r"PRERREQUISITOS.*|CORREQUISITOS.*)(\s+(TEM[ÁA]TICA|SUBTEM[ÁA]TICA|CORREQUISITOS))*", t))


def _unidades_siac(filas: list[list[str]]) -> list[dict]:
    """Las unidades, sabiendo que **una fila puede traer varias**.

    En Investigación la columna del número dice «3 4» y «11 12»: son unidades que se dictan juntas.
    Y cuando el número va combinado en vertical («11 12» abarcando dos filas), la segunda fila llega
    sin número y le toca el siguiente pendiente. Con `fullmatch(r"\\d+")` —lo evidente— se perdían
    cinco de las doce unidades de Investigación sin decir nada.
    """
    unidades: list[dict] = []
    pendientes: list[int] = []
    for fila in filas:
        nums = [int(x) for x in re.findall(r"\d+", fila[0])] if fila else []
        tem = (fila[1] if len(fila) > 1 else "").strip()
        sub = (fila[2] if len(fila) > 2 else "").strip()
        if nums:
            if pendientes and unidades:      # sobraron de la fila anterior: esa fila cubría varias
                unidades[-1]["tambien"] = pendientes
            pendientes = nums
        if not tem and not sub:
            continue
        unidades.append({"n": pendientes.pop(0) if pendientes else None,
                         "tematica": tem, "subtematica": sub})
    if pendientes and unidades:
        unidades[-1]["tambien"] = pendientes
    return unidades


def _unidades_didacticas(doc) -> list[dict]:
    """La tabla «Unidad didáctica · ¿Qué debe enseñar el docente? · Producto esperado».

    Nació con el sílabo de especialización, pero **TG2 la usa dentro del armazón SIAC de pregrado**:
    el documento tiene las 19 tablas del formato 1 y, en lugar de «UNIDADES DE CONOCIMIENTO», esta
    otra. Por eso vive aparte y la llaman los dos lectores, en vez de duplicarse.
    """
    for tabla in doc.tables:
        filas = _filas(tabla)
        if not filas or "unidad didáctica" not in " ".join(filas[0]).lower():
            continue
        salida = []
        for fila in filas[1:]:
            titulo = (fila[0] if fila else "").strip()
            if not titulo:
                continue
            m = re.match(r"^\s*(\d+)[.\-)]?\s*(.*)$", titulo)
            salida.append({
                "n": int(m.group(1)) if m else None,
                "tematica": (m.group(2) if m else titulo).strip(),
                "subtematica": " · ".join(x.strip() for x in fila[1:] if x.strip()),
            })
        return salida
    return []


def _leer_siac(doc, d: dict) -> None:
    crudo: dict[str, list[list[str]]] = {}
    actual = None
    for tabla in doc.tables:
        for fila in _filas(tabla):
            sec = _seccion_de(" ".join(fila), SECCIONES_SIAC)
            if sec:
                actual = sec
                crudo.setdefault(actual, [])
                continue
            if actual and actual != "_tramite" and any(fila) and not _es_rotulo(fila):
                crudo[actual].append(fila)

    for clave in TEXTO:
        if crudo.get(clave):
            d[clave] = " ".join(" ".join(f) for f in crudo[clave]).strip()
    d["evaluacion"] = " · ".join(" ".join(f) for f in crudo.get("evaluacion", [])).strip()
    d["unidades"] = _unidades_siac(crudo.get("unidades", []))
    if not d["unidades"]:                     # TG2: armazón SIAC con tabla de unidades didácticas
        d["unidades"] = _unidades_didacticas(doc)
    for fila in crudo.get("elementos", []):
        t = " ".join(fila[1:]) if re.fullmatch(r"\d+", fila[0].strip()) else " ".join(fila)
        if t.strip():
            d["elementos"].append(t.strip())
    for fila in crudo.get("bibliografia", []):
        ref = " ".join(x for x in fila if x.strip() not in {"●", "•", "-"}).strip()
        if len(ref) > 12:
            d["bibliografia"].append(ref)


def _leer_especializacion(doc, d: dict) -> None:
    partes: dict[str, list[str]] = {}
    actual = None
    for p in doc.paragraphs:
        t = " ".join(p.text.split())
        if not t:
            continue
        sec = _seccion_de(t, SECCIONES_ESP)
        if sec:
            actual = sec
            partes.setdefault(actual, [])
            continue
        if actual and actual != "_tramite":
            partes[actual].append(t)

    for clave in TEXTO:
        if partes.get(clave):
            d[clave] = " ".join(partes[clave]).strip()
    d["resultado"] = d["resultado"] or d["competencia"]
    d["bibliografia"] = [r for r in partes.get("bibliografia", []) if len(r) > 12]

    d["unidades"].extend(_unidades_didacticas(doc))


def ficha(curso: str) -> dict:
    """Lo que el Syllabus dice de la asignatura, en piezas usables para redactar comentarios.

    Siempre devuelve las mismas claves, aunque no haya Syllabus: quien lo use no tiene que preguntar.
    """
    c = SC.COURSES[curso]
    d: dict = {
        "curso": curso, "asignatura": c["titulo_largo"], "codigo": c.get("codigo"),
        "fuente": c.get("fuente"), "nota": c.get("nota_syllabus"), "archivo": None, "formato": None,
        "descripcion": "", "proposito": "", "competencia": "", "resultado": "",
        "mecanismos": "", "evaluacion": "", "estrategia": "",
        "elementos": [], "unidades": [], "bibliografia": [], "advertencias": [],
    }
    ruta = archivo(curso)
    if ruta is None:
        d["advertencias"].append(d["nota"] or "La CUN no entregó el Syllabus de esta asignatura.")
        return d
    d["archivo"] = str(ruta.relative_to(REPO))

    from docx import Document

    doc = Document(str(ruta))
    todo_tablas = " ".join(" ".join(f[0] for f in _filas(t)[:2] if f) for t in doc.tables).upper()
    if "DATOS GENERALES DE LA UNIDAD CURRICULAR" in todo_tablas:
        d["formato"] = "siac"
        _leer_siac(doc, d)
    elif "UNIDAD DIDÁCTICA" in todo_tablas.replace("DIDACTICA", "DIDÁCTICA"):
        d["formato"] = "especializacion"
        _leer_especializacion(doc, d)
    else:
        d["advertencias"].append(
            f"{ruta.name} no tiene ninguno de los dos formatos conocidos (SIAC ni especialización): "
            "no leí nada de él. Ábrelo a mano.")
        return d

    # El Syllabus numera las entregas ACA 1/2/3; el aula no. Quien comente tiene que saberlo, porque
    # el estudiante ve los nombres del aula y no los del Syllabus.
    if re.search(r"\bACA\s*[123]\b", d["evaluacion"] + " " + d["mecanismos"], re.I) or \
            d["formato"] == "especializacion":
        d["advertencias"].append(
            "La numeración de entregas del Syllabus (ACA 1/2/3) NO es la del aula. Los nombres y los "
            "pesos vigentes salen de `config/cursos/fechas_entrega_aca.py`.")
    faltan = [n for n in range(1, max((u["n"] or 0) for u in d["unidades"]) + 1)
              if d["unidades"] and n not in {u["n"] for u in d["unidades"]}
              and n not in {x for u in d["unidades"] for x in u.get("tambien", [])}]
    if faltan:
        d["advertencias"].append(
            f"El Syllabus se salta la unidad {', '.join(map(str, faltan))}: la numeración es del "
            "documento oficial, no un error de lectura.")
    return d


def resumen(curso: str, *, ancho: int = 96) -> list[str]:
    """La ficha en líneas listas para imprimir, que es como la lee quien redacta el plan."""
    f = ficha(curso)
    L = [f"Asignatura : {f['asignatura']}" + (f" · {f['codigo']}" if f["codigo"] else "")]
    if not f["archivo"]:
        L += ["Syllabus   : NO HAY Syllabus en el repositorio.",
              f"             {f['advertencias'][0]}",
              f"             Lo que hay: {f['fuente']}.",
              "             No te inventes competencias ni unidades: comenta con la guía del ACA."]
        return L
    L.append(f"Syllabus   : {f['archivo']}  ({f['formato']})")
    for etq, clave in (("Competencia", "competencia"), ("Resultado", "resultado")):
        if f[clave]:
            L.append(f"{etq:<11}: {f[clave][:ancho * 3]}")
    if f["elementos"]:
        L.append(f"Elementos de competencia ({len(f['elementos'])}):")
        L += [f"  {i}. {t[:ancho]}" for i, t in enumerate(f["elementos"], 1)]
    if f["unidades"]:
        L.append(f"Unidades ({len(f['unidades'])}):")
        for u in f["unidades"]:
            n = f"U{u['n']}" + (f"+{'+'.join(map(str, u['tambien']))}" if u.get("tambien") else "")
            L.append(f"  {n:<7} {u['tematica'][:ancho]}")
    if f["mecanismos"]:
        L.append(f"Evaluación : {f['mecanismos'][:ancho * 3]}")
    for x in f["advertencias"]:
        L.append(f"AVISO      : {x}")
    return L


if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(description="Lo que el Syllabus dice de la asignatura.")
    ap.add_argument("curso", nargs="?", choices=sorted(SC.COURSES), help="sin curso: los lista todos")
    a = ap.parse_args()

    for curso in ([a.curso] if a.curso else sorted(SC.COURSES)):
        print(f"\n{'─' * 100}\n{curso}")
        for linea in resumen(curso):
            print("  " + linea)
