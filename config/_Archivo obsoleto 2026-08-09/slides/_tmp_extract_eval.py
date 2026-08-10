# -*- coding: utf-8 -*-
from docx import Document
from pathlib import Path
import re

files = {
    "P1": Path(
        r"G:\Mi unidad\Trabajos\Empleo\CUN\Cursos\Especializacion\PROYECTO I"
        r"\Especializacion_En_Inteligencia_Artificial_Proyecto_I_ESP329.docx"
    ),
    "INV": Path(
        r"G:\Mi unidad\Trabajos\Empleo\CUN\Cursos\Pregrado"
        r"\INVESTIGACION CIENCIA Y TECNOLOGIA PARA ESCUELA DE INGENIERIAS"
        r"\INVESTIGACION CIENCIA Y TECNOLOGIA PARA ESCUELA DE INGENIERIAS EI005_PRES.docx"
    ),
    "CRE": Path(
        r"G:\Mi unidad\Trabajos\Empleo\CUN\Cursos\Pregrado"
        r"\CREATIVIDAD Y PENSAMIENTO INNOVADOR PARA ESCUELA DE INGENIERIAS"
        r"\CREATIVIDAD Y PENSAMIENTO INNOVADOR PARA ESCUELA DE INGENIERIAS EI004_VIR.docx"
    ),
    "TG3": Path(
        r"G:\Mi unidad\Trabajos\Empleo\CUN\Cursos\Pregrado"
        r"\TRABAJO DE GRADO 3 - MODELOS DE INNOVACION INGENIERIA DE SISTEMAS"
        r"\TRABAJO DE GRADO 3-MDI_INGENIERIA DE SISTEMAS_94532_PRES_VIR.docx"
    ),
}

kw = re.compile(
    r"evalu|aca\b|corte|parcial|entreg|producto|art[ií]culo|rubric|porcent|%|"
    r"EV0|examen|sustenta|propuesta|avance|competenc|actividad",
    re.I,
)

out_dir = Path(r"G:\Mi unidad\Trabajos\Empleo\CUN\Cursos\config\slides\_tmp_eval_extract")
out_dir.mkdir(exist_ok=True)

for key, path in files.items():
    lines = []
    lines.append("=" * 80)
    lines.append(f"{key} {path.name}")
    lines.append("=" * 80)
    doc = Document(str(path))
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if not t:
            continue
        if kw.search(t):
            lines.append(f"P{i}: {t}")
    lines.append("--- TABLES ---")
    for ti, table in enumerate(doc.tables):
        rows_txt = []
        hit = False
        for row in table.rows:
            cells = [" ".join(c.text.split()) for c in row.cells]
            line = " | ".join(cells)
            if kw.search(line):
                hit = True
            rows_txt.append(line)
        if hit:
            lines.append(f"\n[TABLE {ti}]")
            for line in rows_txt:
                if line.strip():
                    lines.append(line)
    text = "\n".join(lines)
    (out_dir / f"{key}.txt").write_text(text, encoding="utf-8")
    print("OK", key, "chars", len(text))
